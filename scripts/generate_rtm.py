from __future__ import annotations

import csv
import re
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PRD = Path("/Users/kishore/Antigravity_Development/Important_Documents/physical_world_interaction_master_spec.docx")
OUT_MD = ROOT / "docs" / "requirements-traceability-matrix.md"
OUT_CSV = ROOT / "docs" / "requirements-traceability-matrix.csv"
OUT_XLSX = ROOT / "docs" / "requirements-traceability-matrix.xlsx"


REFERENCE_RULES = [
    {
        "name": "offline_first_runtime",
        "keywords": [
            "offline",
            "local queue",
            "queue",
            "sync",
            "idempotent",
            "duplicate tap",
            "duplicate prevention",
            "local_event_id",
            "queue_sequence_number",
            "replay",
            "stored locally",
            "locally first",
            "tap must be stored",
            "no duplicate inserts"
        ],
        "status": "Implemented",
        "refs": [
            "packages/runtime/src/queue-store.mjs",
            "packages/runtime/test/queue-store.test.mjs",
            "apps/api/src/routes.mjs",
            "apps/api/test/foundation.test.mjs",
            "apps/api/test/iot-integration.test.mjs"
        ],
        "notes": "Offline queue primitives, idempotent sync, duplicate handling, and replay tests exist. Full Android host remains deferred."
    },
    {
        "name": "device_ops",
        "keywords": ["device", "heartbeat", "incident", "reader", "battery", "diagnostics", "fleet", "assignment", "config", "kiosk"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/src/iot",
            "apps/api/test/iot-integration.test.mjs",
            "apps/web/kiosk.html",
            "apps/web/organizer.html"
        ],
        "notes": "Device config, heartbeat, incidents, credentials, fleet views, IoT health, and go-live readiness are implemented for the pilot build."
    },
    {
        "name": "consent_privacy_masking",
        "keywords": ["consent", "pii", "personal data", "masked", "masking", "privacy", "revoke", "own connections", "attendee"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/masking.mjs",
            "apps/api/src/policy.mjs",
            "apps/api/src/routes.mjs",
            "apps/api/src/session-tokens.mjs",
            "docs/spec-closure-pack/masking-matrix.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Consent-aware release, attendee session tokens, masking by role, sponsor PII gating, revoke, DSR, and audit coverage are implemented."
    },
    {
        "name": "auth_rbac_iam",
        "keywords": ["auth", "oidc", "jwt", "role", "permission", "rbac", "tenant", "event scope", "scope", "login", "user", "admin"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/app.mjs",
            "apps/api/src/auth/oidc.mjs",
            "apps/api/src/auth/principals.mjs",
            "apps/api/src/access-control.mjs",
            "apps/api/src/policy.mjs",
            "apps/web/admin.html",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "OIDC, seed-demo tokens, IAM lifecycle, role/scope enforcement, route access-control matrix, and admin UI are implemented."
    },
    {
        "name": "api_middleware_contract",
        "keywords": [
            "requestidmiddleware",
            "request id",
            "request_id",
            "transportsecuritycheckmiddleware",
            "tenantresolutionmiddleware",
            "resource resolution",
            "resourceresolutionmiddleware",
            "rolescopemiddleware",
            "policyenginemiddleware",
            "validationmiddleware",
            "endpoint handler",
            "responsemaskingmiddleware",
            "auditmiddleware",
            "metricsmiddleware",
            "protected request",
            "middleware order",
            "error contract",
            "all errors must return",
            "all uuids validated",
            "tap_type only"
        ],
        "status": "Implemented",
        "refs": [
            "apps/api/src/app.mjs",
            "apps/api/src/routes.mjs",
            "apps/api/src/server.mjs",
            "apps/api/src/masking.mjs",
            "apps/api/src/policy.mjs",
            "apps/api/test/foundation.test.mjs",
            "packages/runtime/test/http-handler.test.mjs"
        ],
        "notes": "The modular app pipeline covers request IDs, auth, tenant/resource resolution, role/scope checks, policy checks, validation, masking, auditing, metrics, and safe error responses."
    },
    {
        "name": "exports_audit_break_glass",
        "keywords": [
            "export",
            "approval",
            "download",
            "audit",
            "break-glass",
            "break glass",
            "sensitive action",
            "expiry",
            "signed",
            "report freeze",
            "freeze approved",
            "requester cannot approve",
            "cannot approve their own",
            "expired access",
            "maximum duration"
        ],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/src/policy.mjs",
            "apps/web/organizer.html",
            "apps/web/admin.html",
            "docs/spec-closure-pack/export-rules.md",
            "docs/spec-closure-pack/break-glass-sop.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Export requests, approvals, downloads, audit logs, break-glass approvals, revoke, and cross-linked investigation flows are implemented."
    },
    {
        "name": "sponsor_dashboard",
        "keywords": ["sponsor", "impression", "ctr", "campaign", "audience insight", "snapshot"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/web/sponsor.html",
            "docs/spec-closure-pack/metric-versioning-policy.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Sponsor aggregate dashboard, consent-limited PII, report snapshots, export/download, and report freeze workflows are implemented."
    },
    {
        "name": "vendor_dashboard_crm",
        "keywords": ["vendor", "lead", "hot", "warm", "cold", "notes", "crm", "salesforce", "hubspot", "zoho"],
        "status": "Partially Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/src/crm/pilot-crm.mjs",
            "apps/web/vendor.html",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Vendor lead inbox/detail, classification, notes, export requests, and pilot CRM sync exist. Provider-specific Salesforce/HubSpot/Zoho connectors remain future integration work."
    },
    {
        "name": "organizer_operations",
        "keywords": ["organizer", "event overview", "data control", "traffic", "analytics", "incident", "runbook", "go-live", "signoff"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/web/organizer.html",
            "deploy/staging/PILOT_GO_LIVE_RUNBOOK.md",
            "deploy/staging/JOINT_PILOT_SIGNOFF_EXECUTION.md",
            "apps/api/test/foundation.test.mjs",
            "apps/api/test/iot-integration.test.mjs"
        ],
        "notes": "Organizer ops includes overview, fleet, incidents, audit/export/break-glass investigation links, rehearsal, signoff, and go-live execution."
    },
    {
        "name": "compliance_lifecycle",
        "keywords": ["dsr", "data subject", "retention", "delete", "deletion", "anonym", "compliance", "post-event"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/src/crm/pilot-crm.mjs",
            "apps/web/organizer.html",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "DSR, retention preview/apply, anonymization, downstream deletion propagation, CRM cleanup, and compliance reporting are implemented."
    },
    {
        "name": "database_persistence",
        "keywords": [
            "postgres",
            "database",
            "table",
            "schema",
            "rls",
            "row level",
            "tenant isolation",
            "redis",
            "object storage",
            "event_data_policies",
            "event data policy",
            "created_at on tap_events",
            "cloud_received_at",
            "field definitions",
            "allowed values"
        ],
        "status": "Partially Implemented",
        "refs": [
            "apps/api/migrations",
            "apps/api/src/repositories/postgres.mjs",
            "apps/api/src/db/postgres.mjs",
            "docs/spec-closure-pack/schema-v1.sql",
            "apps/api/test/postgres.integration.test.mjs"
        ],
        "notes": "Postgres schema, migrations, runtime role, and RLS hardening exist. Redis/object storage are not required for the local pilot scaffold and remain production infrastructure work."
    },
    {
        "name": "security_deployment_pentest",
        "keywords": ["security", "tls", "rate limiting", "cors", "hsts", "csp", "penetration", "pentest", "production", "deployment", "readiness", "go/no-go"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/security-hardening.mjs",
            "apps/api/src/deployment-readiness.mjs",
            "apps/api/src/server.mjs",
            "deploy/production",
            "docs/spec-closure-pack/security-hardening-and-pentest-readiness.md",
            "docs/spec-closure-pack/external-pentest-support.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Security headers, rate limiting, readiness, no-store responses, sanitized errors, pen-test findings, and production checklists are implemented."
    },
    {
        "name": "final_launch_package",
        "keywords": ["final", "launch", "freeze window", "change control", "approver", "approval matrix"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/migrations/018_final_launch_approvals.sql",
            "docs/spec-closure-pack/final-go-live-package.md",
            "deploy/production/FINAL_GO_LIVE_CHECKLIST.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Final go-live package, four-role approvals, launch export, and post-launch monitoring checklist are implemented."
    },
    {
        "name": "notifications_messaging",
        "keywords": ["notification", "notification_attempt", "outbound message", "followup", "sms", "whatsapp", "marketing notification"],
        "status": "Deferred/Gap",
        "refs": [
            "docs/spec-closure-pack/schema-v1.sql",
            "README.md"
        ],
        "notes": "Notification/follow-up tables are documented in the PRD/spec pack but full outbound messaging workers are outside the current pilot implementation."
    },
    {
        "name": "wallet_short_links",
        "keywords": ["wallet", "wallet_pass", "short link", "short_links"],
        "status": "Deferred/Gap",
        "refs": [
            "docs/spec-closure-pack/pilot-scope-split.md",
            "README.md"
        ],
        "notes": "Wallet passes and general short-link platform are deferred; attendee session tokens and export download routes cover pilot-critical signed access."
    },
    {
        "name": "branding_assets",
        "keywords": ["branding", "logo", "creative", "asset", "idle message", "sponsor panel", "cache"],
        "status": "Partially Implemented",
        "refs": [
            "apps/web/kiosk.html",
            "apps/web/sponsor.html",
            "docs/spec-closure-pack/pilot-scope-split.md",
            "apps/api/src/iot/mock-fixtures.mjs"
        ],
        "notes": "Static/demo UI and fixtures cover branding display expectations. Versioned remote branding asset publishing is not fully implemented."
    },
    {
        "name": "commercial_partner_model",
        "keywords": ["commercial", "sales", "partner", "payout", "billing", "pricing", "proposal", "pipeline", "closed won", "closed lost"],
        "status": "Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/src/access-control.mjs",
            "apps/api/migrations/019_commercial_partner_governance.sql",
            "docs/spec-closure-pack/commercial-partner-governance.md",
            "README.md",
            "apps/api/test/foundation.test.mjs"
        ],
        "notes": "Mandatory commercial partner governance now covers fixed partner types, status-only partner updates, explicit platform-access provisioning, fixed sales stages, required next action/date, ROI-led positioning, founder/product-owner exception approvals, and partner payout tracking through approval/payment after client payment receipt."
    },
    {
        "name": "webhook_delivery",
        "keywords": ["webhook", "webhook_subscriptions", "webhook_deliveries", "target_url"],
        "status": "Partially Implemented",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/test/foundation.test.mjs",
            "docs/spec-closure-pack/schema-v1.sql"
        ],
        "notes": "Downstream webhook deletion propagation exists for compliance flows. General outbound webhook subscriptions and delivery-attempt tracking remain broader integration work."
    },
    {
        "name": "ai_enrichment",
        "keywords": ["ai", "enrichment", "agent orchestrator", "third-party enrichment"],
        "status": "Implemented by Exclusion",
        "refs": [
            "apps/api/src/routes.mjs",
            "apps/api/test/foundation.test.mjs",
            "README.md"
        ],
        "notes": "No AI or third-party enrichment runs in the critical tap path. Optional enrichment workers are not implemented in the pilot build."
    },
    {
        "name": "android_host",
        "keywords": ["android", "locked pwa", "nfc adapter", "kiosk mode", "reader connectivity", "qr fallback"],
        "status": "Partially Implemented",
        "refs": [
            "apps/web/kiosk.html",
            "docs/spec-closure-pack/runtime-adr.md",
            "docs/spec-closure-pack/kiosk-state-machine.md",
            "README.md"
        ],
        "notes": "Web kiosk and runtime state model exist. Native Android host and physical NFC adapter integration remain deferred."
    },
    {
        "name": "ui_states",
        "keywords": ["loading", "empty", "error state", "success", "failure feedback", "screen"],
        "status": "Partially Implemented",
        "refs": [
            "apps/web/attendee.html",
            "apps/web/vendor.html",
            "apps/web/sponsor.html",
            "apps/web/organizer.html",
            "apps/web/admin.html",
            "apps/web/test/phase3-ui.e2e.test.mjs"
        ],
        "notes": "Static web shells include loading/empty/error feedback in major flows. A formal design-system-wide state audit remains useful before production UI polish."
    },
]


REQUIREMENT_TERMS = re.compile(
    r"\b(must|shall|required|mandatory|never|always|only|cannot|can not|no |every |all |should|allowed|forbidden|gates|requires|required)\b",
    re.IGNORECASE,
)


STATUS_RANK = {
    "Deferred/Gap": 5,
    "Partially Implemented": 4,
    "Implemented by Exclusion": 2,
    "Implemented": 1,
}


STATUS_OVERRIDES = {
    "RTM-0014": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "The public leaderboard endpoint returns only aggregate stall rankings and generalized latest-connection ticker text; personal names, emails, phone numbers, and exact attendee company names are excluded.",
    },
    "RTM-0017": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, sponsor_dashboard, exports_audit_break_glass",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/policy.mjs; apps/api/src/masking.mjs; apps/web/sponsor.html",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Sponsor users receive aggregate sponsor dashboard metrics with explicit no-PII privacy metadata, and sponsor lead exports include only vendor_and_sponsor opted-in leads when sponsor_pii_enabled remains true.",
    },
    "RTM-0030": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking",
        "code_references": "apps/api/src/masking.mjs; apps/api/src/routes.mjs; apps/api/src/policy.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Lead list/detail responses now carry explicit privacy reasons and mask name, company, title, email, and phone whenever consent, role, event policy, or break-glass state does not permit display.",
    },
    "RTM-0088": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, commercial_partner_model",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/crm/pilot-crm.mjs; apps/api/migrations/019_commercial_partner_governance.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Commercial deal governance and pilot CRM sync now use the fixed sales pipeline stages: Lead Added, Contacted, Replied, Call Scheduled, Demo Done, Proposal Sent, Negotiation, Closed Won, and Closed Lost.",
    },
    "RTM-0168": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Automated coverage verifies the public leaderboard payload contains no attendee personal data even when the underlying interaction has name, email, phone, and exact company profile data.",
    },
    "RTM-0203": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, sponsor_dashboard, exports_audit_break_glass",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/policy.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Sponsor lead exports are rebuilt at download time and include only interactions whose current consent_status is vendor_and_sponsor; vendor-only, pending, declined, and revoked leads are excluded.",
    },
    "RTM-0254": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, sponsor_dashboard, exports_audit_break_glass",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/policy.mjs; apps/api/src/masking.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "When sponsor_pii_enabled is false, sponsor lead export requests are denied, sponsor dashboard output remains aggregate-only, and masked lead responses explain the event-policy privacy reason.",
    },
    "RTM-0299": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, auth_rbac_iam, vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/masking.mjs; apps/api/src/policy.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Lead detail access is role-scoped and consent-aware for vendor, organizer, and platform users; platform unmasking requires active scoped break-glass, and sponsor users remain on aggregate/export paths rather than stall lead detail.",
    },
    "RTM-0301": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/masking.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "When vendor consent is absent or later revoked, lead list/detail returns the anonymous masked view with privacy.reason=vendor_consent_required and CRM eligibility blocked by consent.",
    },
    "RTM-0322": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, sponsor_dashboard",
        "code_references": "apps/api/src/routes.mjs; apps/web/sponsor.html",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Sponsor metrics are aggregate-only with privacy.personal_data_included=false; sponsor PII can only flow through controlled sponsor lead export after policy and consent checks pass.",
    },
    "RTM-0325": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/policy.mjs; apps/api/src/crm/pilot-crm.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "CRM push is blocked unless current consent_status is vendor_only or vendor_and_sponsor and event policy allows CRM; revocation moves eligibility back to blocked_by_consent and CRM sync returns 409.",
    },
    "RTM-0089": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, commercial_partner_model",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/crm/pilot-crm.mjs; apps/api/migrations/019_commercial_partner_governance.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Every commercial deal and pilot CRM upsert now carries stage, next_action, and next_action_at; create/update validation rejects missing or blank required CRM hygiene fields.",
    },
    "RTM-0096": {
        "coverage_status": "Implemented",
        "implementation_area": "database_persistence, organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/api/migrations/001_init.sql; apps/api/migrations/021_leaderboard_snapshots.sql; docs/spec-closure-pack/schema-v1.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "leaderboard_snapshots is implemented as a persisted snapshot table/repository plus organizer capture/list APIs for historical leaderboard replay and reporting.",
    },
    "RTM-0306": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/masking.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs; apps/web/test/phase3-ui.e2e.test.mjs",
        "notes": "Vendor lead inbox now returns the required column contract and the browser inbox renders timestamp, attendee, company, title, score, consent status, next action, CRM state, and notes-aware detail with consent masking.",
    },
    "RTM-0307": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs; apps/web/test/phase3-ui.e2e.test.mjs",
        "notes": "Vendor lead inbox now supports bounded limit/offset pagination with total, has_more, next_offset, browser page size, and previous/next controls.",
    },
    "RTM-0323": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "GET /events/{id}/leaderboard is implemented as a public no-PII endpoint with aggregate rankings and generalized connection ticker entries only.",
    },
    "RTM-0348": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/migrations/001_init.sql; apps/api/migrations/020_lead_score_history.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Lead scoring is constrained to hot, warm, and cold at API validation and database-check levels, with score-history events using the same fixed values.",
    },
    "RTM-0369": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Public leaderboard output is generated from aggregate counts and generalized ticker descriptors only; automated tests verify no personal profile data appears in the payload.",
    },
    "RTM-0415": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, database_persistence",
        "code_references": "apps/api/migrations/020_lead_score_history.sql; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/api/src/routes.mjs; docs/spec-closure-pack/schema-v1.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Lead score history is now persisted through the lead_scores table/repository and returned on consent-aware lead detail responses for analytics and auditability.",
    },
    "RTM-0416": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, database_persistence",
        "code_references": "apps/api/migrations/020_lead_score_history.sql; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/api/src/routes.mjs; docs/spec-closure-pack/schema-v1.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "The starter latest-score field remains for fast reads, and every score change now appends a lead_scores history row with previous score, new score, actor, reason, and timestamp.",
    },
    "RTM-0438": {
        "coverage_status": "Implemented",
        "implementation_area": "public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Latest connection ticker text now uses the allowed generalized form, for example: Someone from a large enterprise connected with [Stall Name].",
    },
    "RTM-0439": {
        "coverage_status": "Implemented",
        "implementation_area": "public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Public leaderboard ticker entries never include attendee personal names; regression tests assert seeded names are absent from serialized output.",
    },
    "RTM-0440": {
        "coverage_status": "Implemented",
        "implementation_area": "public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Public leaderboard ticker entries never include attendee email addresses; regression tests assert seeded emails are absent from serialized output.",
    },
    "RTM-0441": {
        "coverage_status": "Implemented",
        "implementation_area": "public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Public leaderboard ticker entries use generalized company descriptors by default and suppress exact attendee company names unless a future policy explicitly enables exact company display with legal basis.",
    },
    "RTM-0509": {
        "coverage_status": "Implemented",
        "implementation_area": "database_persistence, organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/api/migrations/021_leaderboard_snapshots.sql; docs/spec-closure-pack/schema-v1.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Organizer leaderboard snapshots now provide the historical analytics/replay store required by the PRD, including snapshot payloads with rankings and latest connection ticker state.",
    },
    "RTM-0510": {
        "coverage_status": "Implemented",
        "implementation_area": "database_persistence, organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/api/migrations/021_leaderboard_snapshots.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Because public leaderboard history/reporting is in scope, the conditional leaderboard_snapshots requirement is implemented through persisted snapshots and organizer replay APIs.",
    },
    "RTM-0511": {
        "coverage_status": "Implemented",
        "implementation_area": "organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/migrations/021_leaderboard_snapshots.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Leaderboard snapshot capture records snapshot_interval_minutes=5 and enforces the five-minute cadence window unless an organizer explicitly forces a capture.",
    },
    "RTM-0512": {
        "coverage_status": "Implemented",
        "implementation_area": "organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Each leaderboard snapshot payload records the exact ranking formula: count all event interactions per stall, sort by connection_count descending, then stall_name ascending.",
    },
    "RTM-0513": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/api/src/masking.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Leaderboard snapshots reuse the public no-PII leaderboard payload and regression tests assert attendee names, emails, phone numbers, and exact company profile values are absent.",
    },
    "RTM-0435": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Vendor dashboard metrics now count all scoped stall interactions in the selected period through /stalls/{stall_id}/dashboard-metrics and surface Total Taps in the vendor UI.",
    },
    "RTM-0436": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Response Rate is explicitly defined and implemented as distinct CRM-pushed or follow-up-sent leads divided by distinct vendor-consented leads, returning 0 when the denominator is 0.",
    },
    "RTM-0791": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, vendor_dashboard_crm, ui_states",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/masking.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs; apps/web/test/phase3-ui.e2e.test.mjs",
        "notes": "Vendor Lead Inbox is chronological, paginated, filterable by score, consent, and CRM eligibility, and shows timestamp, attendee, company, title, score, consent status, next action, and CRM state while preserving anonymous masking when vendor consent is absent.",
    },
    "RTM-0793": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, ui_states",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs; apps/web/test/phase3-ui.e2e.test.mjs",
        "notes": "Vendor Lead Detail supports inline Hot/Warm/Cold scoring and note capture, and now displays score history beside interaction notes.",
    },
    "RTM-0792": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/masking.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs; apps/web/test/phase3-ui.e2e.test.mjs",
        "notes": "Vendor lead detail includes profile, interaction state, notes, score history, follow-up/CRM state, and hides PII with explicit privacy metadata whenever consent is unavailable.",
    },
    "RTM-0789": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, compliance_lifecycle",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/compliance/post-event-lifecycle.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "The signed attendee session shows the attendee's stored connection detail and exposes self-service controls to revoke consent, request an access export, or file a delete request into the organizer DSR queue.",
    },
    "RTM-0790": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, compliance_lifecycle",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/compliance/post-event-lifecycle.mjs; apps/api/migrations/022_attendee_self_service_dsr.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Attendee privacy controls support sponsor opt-out while retaining vendor consent, full consent revocation, access-export DSR creation, and delete DSR creation from the signed attendee session.",
    },
    "RTM-0817": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, sponsor_dashboard, vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/policy.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "POST /consents/revoke now revokes vendor and/or sponsor release rights prospectively; sponsor opt-out removes future sponsor exports while keeping vendor access/CRM eligible, and full revoke masks detail plus blocks CRM and exports.",
    },
    "RTM-0822": {
        "coverage_status": "Implemented",
        "implementation_area": "consent_privacy_masking, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "GET /events/{id}/leaderboard now returns a leaderboard dataset with aggregate rankings and no personal data.",
    },
    "RTM-0846": {
        "coverage_status": "Implemented",
        "implementation_area": "auth_rbac_iam, database_persistence, organizer_operations, public_leaderboard",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/api/src/repositories/memory.mjs; apps/api/src/repositories/postgres.mjs; apps/api/migrations/001_init.sql; apps/api/migrations/021_leaderboard_snapshots.sql; docs/spec-closure-pack/schema-v1.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "leaderboard_snapshots is now an authoritative historical leaderboard-state scope with organizer-only create/list routes and tenant-scoped persistence.",
    },
    "RTM-0884": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/access-control.mjs; apps/web/vendor.html",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "Response Rate uses the fixed PRD formula: (distinct CRM pushed or followup sent) / distinct vendor-consented leads, with 0 returned when denominator is 0.",
    },
    "RTM-0748": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, commercial_partner_model",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/crm/pilot-crm.mjs; apps/api/migrations/019_commercial_partner_governance.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "CRM hygiene is enforced for commercial sales records and pilot CRM lead upserts: stage, next action, and next action date are always present and validated.",
    },
    "RTM-0749": {
        "coverage_status": "Implemented",
        "implementation_area": "vendor_dashboard_crm, commercial_partner_model",
        "code_references": "apps/api/src/routes.mjs; apps/api/src/crm/pilot-crm.mjs; apps/api/migrations/019_commercial_partner_governance.sql",
        "test_references": "apps/api/test/foundation.test.mjs",
        "notes": "No synced lead or commercial deal may remain without a next action; pilot CRM assigns classification-aware follow-up actions and deal APIs reject blank next_action values.",
    },
}


def clean(value: str) -> str:
    return " ".join((value or "").split())


def extract_requirements(prd_path: Path) -> list[dict[str, str]]:
    doc = Document(prd_path)
    rows: list[dict[str, str]] = []
    section_stack: list[str] = []
    seen: set[tuple[str, str]] = set()

    def section() -> str:
        return " > ".join(section_stack[-4:]) if section_stack else "Document"

    def add(source: str, req_type: str, text: str) -> None:
        text = clean(text)
        if not text or len(text) < 8:
            return
        key = (section(), text.lower())
        if key in seen:
            return
        seen.add(key)
        rows.append({
            "source": source,
            "section": section(),
            "type": req_type,
            "requirement": text,
        })

    for index, para in enumerate(doc.paragraphs):
        text = clean(para.text)
        if not text:
            continue
        style = para.style.name or ""
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip().split()[0])
            except Exception:
                level = 1
            section_stack = section_stack[: max(level - 1, 0)]
            section_stack.append(text)
            continue
        if "List" in style:
            add(f"paragraph:{index}", style, text)
        elif REQUIREMENT_TERMS.search(text):
            add(f"paragraph:{index}", style, text)

    for table_index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers = [clean(cell.text) for cell in table.rows[0].cells]
        if headers == ["Field", "Value"]:
            continue
        for row_index, row in enumerate(table.rows[1:], start=1):
            values = [clean(cell.text) for cell in row.cells]
            if not any(values):
                continue
            if len(values) == 2:
                text = f"{headers[0]}: {values[0]}; {headers[1]}: {values[1]}"
            else:
                text = "; ".join(
                    f"{header or f'Column {idx + 1}'}: {value}"
                    for idx, (header, value) in enumerate(zip(headers, values))
                    if value
                )
            add(f"table:{table_index}:row:{row_index}", "Table row", text)

    return rows


def classify(row: dict[str, str]) -> dict[str, str]:
    haystack = f"{row['section']} {row['requirement']}".lower()
    matches = []
    for rule in REFERENCE_RULES:
        if any(keyword_matches(haystack, keyword) for keyword in rule["keywords"]):
            matches.append(rule)

    if not matches:
        return {
            "coverage_status": "Deferred/Gap",
            "implementation_area": "manual_review_deferred",
            "code_references": "Deferred by production scope decision",
            "test_references": "No direct automated test reference",
            "notes": "Deferred by product decision: the prior Needs Review category is not part of the immediate production build and will require a later formal scope review before implementation.",
        }

    selected = max(matches, key=lambda item: STATUS_RANK.get(item["status"], 0))
    refs = []
    notes = []
    areas = []
    for match in matches:
        areas.append(match["name"])
        refs.extend(match["refs"])
        notes.append(match["notes"])

    unique_refs = list(dict.fromkeys(refs))
    test_refs = [ref for ref in unique_refs if "/test/" in ref or ref.endswith(".test.mjs")]
    code_refs = [ref for ref in unique_refs if ref not in test_refs]

    return {
        "coverage_status": selected["status"],
        "implementation_area": ", ".join(dict.fromkeys(areas)),
        "code_references": "; ".join(code_refs) or "No direct code reference",
        "test_references": "; ".join(test_refs) or "No direct automated test reference",
        "notes": " ".join(dict.fromkeys(notes)),
    }


def keyword_matches(haystack: str, keyword: str) -> bool:
    normalized = keyword.lower()
    if normalized == "ai":
        return re.search(r"\bai\b", haystack) is not None
    return normalized in haystack


def build_matrix(prd_path: Path) -> list[dict[str, str]]:
    requirements = extract_requirements(prd_path)
    matrix = []
    for index, row in enumerate(requirements, start=1):
        requirement_id = f"RTM-{index:04d}"
        mapped = classify(row)
        mapped.update(STATUS_OVERRIDES.get(requirement_id, {}))
        matrix.append({
            "requirement_id": requirement_id,
            **row,
            **mapped,
        })
    return matrix


def write_csv(rows: list[dict[str, str]]) -> None:
    with OUT_CSV.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_xlsx(rows: list[dict[str, str]]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "RTM"
    headers = list(rows[0].keys())
    sheet.append(headers)
    for row in rows:
        sheet.append([row[key] for key in headers])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    widths = {
        "A": 14,
        "B": 18,
        "C": 38,
        "D": 18,
        "E": 80,
        "F": 24,
        "G": 30,
        "H": 70,
        "I": 60,
        "J": 80,
    }
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    summary = workbook.create_sheet("Summary")
    summary.append(["Coverage status", "Count"])
    counts = Counter(row["coverage_status"] for row in rows)
    for status, count in counts.most_common():
        summary.append([status, count])
    for cell in summary[1]:
        cell.font = Font(bold=True)
    summary.column_dimensions["A"].width = 28
    summary.column_dimensions["B"].width = 12

    workbook.save(OUT_XLSX)


def write_markdown(rows: list[dict[str, str]], prd_path: Path) -> None:
    counts = Counter(row["coverage_status"] for row in rows)
    gap_rows = [row for row in rows if row["coverage_status"] in {"Deferred/Gap", "Partially Implemented", "Needs Review"}]
    preview_rows = rows[:120]

    lines = [
        "# Requirements Traceability Matrix",
        "",
        f"Source PRD: `{prd_path}`",
        "",
        "Generated artifacts:",
        f"- Full CSV matrix: `docs/{OUT_CSV.name}`",
        f"- Filterable Excel matrix: `docs/{OUT_XLSX.name}`",
        "",
        "Method:",
        "- Extracted requirement-like bullets, numbered items, modal requirement paragraphs, and table rows from the PRD.",
        "- Mapped each requirement to implementation areas using keyword-based trace rules plus current build knowledge.",
        "- Marked deferred or partial items explicitly so they are visible during production hardening and acceptance review.",
        "",
        "Coverage summary:",
        "",
        "| Status | Count |",
        "|---|---:|",
    ]
    for status, count in counts.most_common():
        lines.append(f"| {status} | {count} |")

    completed_overrides = [
        (requirement_id, override)
        for requirement_id, override in STATUS_OVERRIDES.items()
        if override.get("coverage_status") == "Implemented"
    ]
    if completed_overrides:
        lines.extend([
            "",
            "Completed build slices:",
            "",
            "| Requirement ID | Area | Completion note |",
            "|---|---|---|",
        ])
        for requirement_id, override in completed_overrides:
            lines.append(
                f"| {requirement_id} | {md_escape(override['implementation_area'])} | {md_escape(override['notes'])} |"
            )

    lines.extend([
        "",
        "High-priority gap register:",
        "",
        "| Requirement ID | Status | Section | Requirement | Notes |",
        "|---|---|---|---|---|",
    ])
    for row in gap_rows[:80]:
        lines.append(
            "| {requirement_id} | {coverage_status} | {section} | {requirement} | {notes} |".format(
                **{key: md_escape(value) for key, value in row.items()}
            )
        )
    if len(gap_rows) > 80:
        lines.append(f"| ... | ... | ... | {len(gap_rows) - 80} additional gap or partial rows are available in the CSV/XLSX matrix. | ... |")

    lines.extend([
        "",
        "RTM preview:",
        "",
        "| Requirement ID | Status | Section | Requirement | Code references | Test references |",
        "|---|---|---|---|---|---|",
    ])
    for row in preview_rows:
        lines.append(
            "| {requirement_id} | {coverage_status} | {section} | {requirement} | {code_references} | {test_references} |".format(
                **{key: md_escape(value) for key, value in row.items()}
            )
        )
    if len(rows) > len(preview_rows):
        lines.append(f"| ... | ... | ... | {len(rows) - len(preview_rows)} additional rows are available in the CSV/XLSX matrix. | ... | ... |")

    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def md_escape(value: str) -> str:
    return clean(str(value)).replace("|", "\\|").replace("\n", " ")


def main() -> None:
    prd_path = DEFAULT_PRD
    rows = build_matrix(prd_path)
    if not rows:
        raise SystemExit("No requirements extracted from PRD.")
    write_csv(rows)
    write_xlsx(rows)
    write_markdown(rows, prd_path)
    counts = Counter(row["coverage_status"] for row in rows)
    print(f"Generated {len(rows)} RTM rows")
    for status, count in counts.most_common():
        print(f"{status}: {count}")
    print(OUT_MD)
    print(OUT_CSV)
    print(OUT_XLSX)


if __name__ == "__main__":
    main()
